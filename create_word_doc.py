#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading(doc, text, level):
    """Добавляет заголовок с правильным форматированием"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = 'Times New Roman'
    heading.style.font.size = Pt(14 if level == 1 else 12)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, align=None):
    """Добавляет параграф с форматированием"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_code_block(doc, code, language=''):
    """Добавляет блок кода"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # Добавляем серый фон для блока кода
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def parse_table_line(line):
    """Парсит строку таблицы markdown"""
    # Удаляем начальные и конечные |
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    # Разделяем по |
    cells = [cell.strip() for cell in line.split('|')]
    return cells

def add_table_from_markdown(doc, lines, start_idx):
    """Добавляет таблицу из markdown"""
    # Парсим заголовок
    header_line = lines[start_idx]
    headers = parse_table_line(header_line)
    
    # Пропускаем разделитель (---)
    separator_idx = start_idx + 1
    
    # Находим конец таблицы
    end_idx = separator_idx + 1
    while end_idx < len(lines) and lines[end_idx].strip().startswith('|'):
        end_idx += 1
    
    # Создаем таблицу
    num_cols = len(headers)
    num_rows = end_idx - separator_idx - 1
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Заполняем заголовок
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    # Заполняем данные
    row_idx = 1
    for i in range(separator_idx + 1, end_idx):
        if i < len(lines):
            cells = parse_table_line(lines[i])
            for j, cell in enumerate(cells):
                if j < num_cols:
                    table.rows[row_idx].cells[j].text = cell
                    table.rows[row_idx].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    table.rows[row_idx].cells[j].paragraphs[0].runs[0].font.size = Pt(11)
            row_idx += 1
    
    return end_idx - 1

def process_markdown_to_word(md_file, output_file):
    """Конвертирует markdown файл в Word документ"""
    doc = Document()
    
    # Настройка стилей по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Обработка блоков кода
        if line.startswith('```'):
            if in_code_block:
                # Конец блока кода
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text, code_language)
                code_block_lines = []
                in_code_block = False
                code_language = ''
            else:
                # Начало блока кода
                in_code_block = True
                code_language = line[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Пропускаем пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Обработка таблиц
        if line.strip().startswith('|') and '---' not in line:
            # Проверяем, есть ли следующая строка с разделителем
            if i + 1 < len(lines) and '---' in lines[i + 1]:
                i = add_table_from_markdown(doc, lines, i)
                i += 1
                continue
        
        # Обработка заголовков
        if line.startswith('# '):
            add_heading(doc, line[2:], 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.startswith('##### '):
            add_heading(doc, line[6:], 5)
        # Обработка разделителей
        elif line.strip() == '---':
            # Добавляем пустую строку вместо разделителя
            doc.add_paragraph()
        # Обработка списков
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Удаляем markdown форматирование
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
        # Обработка нумерованных списков
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
        # Обработка обычного текста
        else:
            # Удаляем markdown форматирование
            text = line
            # Удаляем жирный текст, но оставляем содержимое
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Удаляем курсив
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Удаляем ссылки [текст](url) -> текст
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # Удаляем inline код
            text = re.sub(r'`(.*?)`', r'\1', text)
            # Обработка мест для скриншотов
            if '📸' in text or 'СКРИНШОТ' in text or '[Здесь необходимо' in text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                add_paragraph(doc, text)
        
        i += 1
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"✅ Word документ создан: {output_file}")
    print(f"📄 Файл содержит {len(doc.paragraphs)} параграфов")

if __name__ == '__main__':
    input_file = 'Пояснительная_записка_подробная.md'
    output_file = 'Пояснительная_записка.docx'
    
    try:
        process_markdown_to_word(input_file, output_file)
        print(f"\n✅ Успешно! Документ сохранен как: {output_file}")
    except FileNotFoundError:
        print(f"❌ Ошибка: Файл {input_file} не найден!")
    except Exception as e:
        print(f"❌ Ошибка при создании документа: {e}")
