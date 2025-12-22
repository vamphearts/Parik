#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Устанавливает шрифт для run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

def add_title_page(doc):
    """Добавляет титульный лист"""
    # Добавляем пустые строки для центрирования
    for _ in range(8):
        doc.add_paragraph()
    
    # Название университета
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Федеральное государственное образовательное бюджетное\nучреждение высшего образования\n«Финансовый университет\nпри Правительстве Российской Федерации»\n(Финансовый университет)')
    set_font(run, font_size=14, bold=True)
    
    doc.add_paragraph()
    
    # Департамент
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Департамент анализа данных и машинного обучения')
    set_font(run, font_size=14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Название работы
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Пояснительная записка к курсовой работе по дисциплине\n«Разработка информационных систем»\nна тему:')
    set_font(run, font_size=14)
    
    doc.add_paragraph()
    
    # Тема
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Информационно-справочная система парикмахерской»')
    set_font(run, font_size=14, bold=True)
    
    for _ in range(6):
        doc.add_paragraph()
    
    # Студент
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Выполнил:\nСтудент группы [Номер группы]\n[ФИО студента]')
    set_font(run, font_size=14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Научный руководитель
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Научный руководитель:\n[Ученая степень, должность]\n[ФИО руководителя]')
    set_font(run, font_size=14)
    
    for _ in range(6):
        doc.add_paragraph()
    
    # Город и год
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Москва – 2024')
    set_font(run, font_size=14)
    
    # Разрыв страницы
    doc.add_page_break()

def add_table_of_contents(doc):
    """Добавляет оглавление"""
    # Номер страницы
    p = doc.add_paragraph('2')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.runs[0], font_size=12)
    
    doc.add_paragraph()
    
    # Заголовок
    p = doc.add_paragraph('Оглавление')
    set_font(p.runs[0], font_size=14, bold=True)
    
    doc.add_paragraph()
    
    # Содержание
    toc_items = [
        ('Введение', 4),
        ('1. Описание программы', 7),
        ('1.1. Алгоритмические решения', 7),
        ('1.1.1. Безопасность', 7),
        ('1.1.2. Клиент', 9),
        ('1.2. Описание интерфейса программы', 13),
        ('1.2.1. Навигация и Футер', 13),
        ('1.2.2. Регистрация и Авторизация', 15),
        ('1.2.3. Главная страница', 17),
        ('1.2.4. Управление услугами и мастерами', 19),
        ('1.2.5. Записи и отчеты', 21),
        ('1.3. Архитектура приложения', 22),
        ('1.3.1. Зависимости проекта', 22),
        ('1.3.2. Клиент', 25),
        ('1.3.3. База данных', 26),
        ('2. Структура классов и их назначение в рамках проекта', 28),
        ('2.1. Сервер', 28),
        ('2.1.1. Config', 29),
        ('2.1.2. Controllers', 30),
        ('2.1.3. DTO', 32),
        ('2.1.4. Model', 32),
        ('2.1.5. Repositories', 35),
        ('2.1.6. Security', 36),
        ('2.1.7. Services', 37),
        ('2.2. Клиент', 39),
        ('2.2.1. Template', 39),
        ('2.2.2. Static', 41),
        ('Заключение', 43),
        ('Список использованных источников', 44),
        ('Приложения', 46),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run1 = p.add_run(item)
        set_font(run1, font_size=12)
        
        # Добавляем точки и номер страницы
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15), alignment=1)  # Выравнивание по правому краю
        
        run2 = p.add_run('\t' + str(page))
        set_font(run2, font_size=12)
    
    doc.add_page_break()

def process_content_to_word(content_file, output_file):
    """Создает Word документ из содержимого"""
    doc = Document()
    
    # Настройка полей страницы (как в примере)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # Настройка стилей по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Добавляем титульный лист
    add_title_page(doc)
    
    # Добавляем оглавление
    add_table_of_contents(doc)
    
    # Читаем содержимое
    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    page_num = 4  # Начинаем с 4 страницы (после титульного и оглавления)
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Обработка блоков кода
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                set_font(run, 'Courier New', 10)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Пропускаем пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Обработка заголовков
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], 5)
        # Обработка разделителей
        elif line.strip() == '---':
            doc.add_paragraph()
        # Обработка списков
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            set_font(p.runs[0], font_size=12)
        # Обработка нумерованных списков
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            set_font(p.runs[0], font_size=12)
        # Обработка обычного текста
        else:
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            if '📸' in text or 'СКРИНШОТ' in text or '[Здесь необходимо' in text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_font(run, font_size=12, italic=True)
                run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                p = doc.add_paragraph(text)
                set_font(p.runs[0], font_size=12)
        
        i += 1
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"✅ Word документ создан: {output_file}")

if __name__ == '__main__':
    input_file = 'Пояснительная_записка_подробная.md'
    output_file = 'Пояснительная_записка_точная_копия.docx'
    
    try:
        process_content_to_word(input_file, output_file)
        print(f"\n✅ Успешно! Документ сохранен как: {output_file}")
    except FileNotFoundError:
        print(f"❌ Ошибка: Файл {input_file} не найден!")
    except Exception as e:
        print(f"❌ Ошибка при создании документа: {e}")
        import traceback
        traceback.print_exc()




